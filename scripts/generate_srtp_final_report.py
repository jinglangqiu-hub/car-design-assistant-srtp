from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUT = PROJECT_ROOT / "docs" / "重庆大学SRTP结项报告_汽车外观设计辅助系统.docx"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line=True, line_spacing=1.5, before=0, after=6):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if first_line:
        pf.first_line_indent = Cm(0.74)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFC7D5")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i])
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = OxmlElement("w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), "100")
                node.set(qn("w:type"), "dxa")
                mar.append(node)
            tc_pr.append(mar)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=15, bold=True)
    elif level == 2:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=13, bold=True)
    else:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0, 0, 0)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("重庆大学大学生创新创业训练计划项目")
    set_run_font(r, east_asia="黑体", size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("结项报告")
    set_run_font(r, east_asia="黑体", size=22, bold=True)

    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("基于 StyleGAN2-ADA 的汽车外观设计辅助系统设计与实现")
    set_run_font(r, east_asia="黑体", size=18, bold=True)

    for _ in range(3):
        doc.add_paragraph()
    info = [
        ("项目名称", "汽车外观设计辅助系统"),
        ("项目类型", "SRTP 项目"),
        ("学院/专业", "（请填写学院与专业）"),
        ("项目负责人", "（请填写姓名）"),
        ("项目成员", "（请填写成员姓名）"),
        ("指导教师", "（请填写指导教师）"),
        ("完成时间", "2026 年 6 月"),
    ]
    table = doc.add_table(rows=len(info), cols=2)
    table.autofit = False
    set_table_borders(table)
    for row, (k, v) in zip(table.rows, info):
        set_cell_text(row.cells[0], k, bold=True)
        set_cell_text(row.cells[1], v, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本项目面向汽车外观概念设计早期阶段的辅助生成需求，围绕“可按车型选择生成汽车外观图片”的 SRTP 立项目标，设计并实现了一套基于 StyleGAN2-ADA 的汽车外观设计辅助系统。项目以 CompCars 汽车图像数据为基础，完成了图像清洗、视角筛选、裁剪标准化、车型属性映射和三类训练数据集构建；在云端 RTX 5090 GPU 环境下完成通用基线模型训练、三类车型微调模型训练及追加训练；最终形成跑车、轿车、SUV 三个独立生成模型，并在 Windows 11 桌面程序中实现按车型选择、随机种子控制、truncation 稳定度调节、单张/批量生成、图片保存和参数记录等功能。项目还完成了 PyInstaller 桌面程序打包、Inno Setup 安装包制作、GitHub 源码托管与 Release 资产整理。实验结果表明，系统能够生成具有车辆基本轮廓和类别差异的汽车外观概念图，每类均筛选出 10 张以上可展示样例，基本达到 SRTP 立项阶段对三类汽车外观辅助生成系统的结项要求。"
    )
    add_body(doc, "关键词：汽车外观设计；StyleGAN2-ADA；生成对抗网络；云端训练；桌面应用；SRTP")

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "This project develops a car exterior design assistant based on StyleGAN2-ADA for early-stage automotive concept generation. Using the CompCars dataset, the project builds cleaned and type-specific datasets for sports cars, sedans, and SUVs, trains a general baseline model, and further fine-tunes three type-specific generative models on a cloud RTX 5090 GPU platform. A Windows desktop application is implemented to load the corresponding model according to user selection and generate car exterior images with controllable random seed and truncation parameters. The final system includes trained models, generated samples, a packaged desktop application, an Inno Setup installer, and GitHub-hosted project files. The results show that the system can generate distinguishable vehicle appearances for the three target categories and provides a runnable prototype for SRTP final demonstration."
    )
    add_body(doc, "Keywords: car exterior design; StyleGAN2-ADA; GAN; cloud training; desktop application; SRTP")

    doc.add_page_break()
    add_heading(doc, "目录", 1)
    for item in [
        "1 绪论",
        "2 项目需求与总体方案",
        "3 数据集构建与预处理",
        "4 生成模型方法与训练实现",
        "5 云端迁移与训练过程",
        "6 桌面系统设计与实现",
        "7 实验结果与结项成果",
        "8 问题分析、改进方向与总结",
        "参考文献",
        "附录",
    ]:
        add_body(doc, item)
    doc.add_page_break()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景", 2)
    add_body(
        doc,
        "汽车外观设计通常需要经历需求分析、比例探索、草图绘制、概念评审和细节深化等阶段。在早期概念探索阶段，设计者往往需要快速获得大量具有不同姿态、比例和风格倾向的候选图像，用于启发造型方向和辅助方案讨论。传统人工草图方式具有较强创造性，但生成速度依赖设计经验；常规检索式参考图又容易受到已有车型限制，难以快速形成新的概念组合。近年来，生成式人工智能在图像合成、视觉设计、创意辅助等场景中表现出较强潜力，为汽车外观概念设计提供了新的工具路径。"
    )
    add_body(
        doc,
        "生成对抗网络（Generative Adversarial Network, GAN）通过生成器与判别器的对抗训练学习图像数据分布，能够生成具有较高视觉真实感的图像。StyleGAN 系列模型进一步通过风格映射网络、逐层调制和潜空间控制提升了图像质量与可控性。StyleGAN2-ADA 在小规模数据集场景中引入自适应数据增强策略，适合在有限类别数据上进行微调训练。因此，本项目选择 StyleGAN2-ADA 作为汽车外观图像生成的核心模型，并围绕 SRTP 结项目标构建可运行、可演示、可打包的汽车外观设计辅助系统。"
    )

    add_heading(doc, "1.2 项目目标", 2)
    add_body(
        doc,
        "立项阶段的核心目标是开发一个汽车外观设计辅助系统，使系统能够根据用户选择生成三种类型的汽车外观图，即跑车、轿车和 SUV。围绕该目标，项目需要完成数据集构建、生成模型训练、生成效果评估、桌面端交互程序开发和结项材料整理等工作。最终成果不仅应停留在模型训练文件层面，还应形成普通用户可运行的 Windows 桌面程序，并能通过安装包在其他 Win11 电脑上部署使用。"
    )
    add_bullet(doc, "完成汽车外观数据清洗和标准化处理，形成可用于 StyleGAN2-ADA 训练的数据集。")
    add_bullet(doc, "构建跑车、轿车、SUV 三类训练数据集，并训练对应的三类生成模型。")
    add_bullet(doc, "实现可交互桌面程序，支持车型选择、参数控制、图像生成和保存。")
    add_bullet(doc, "完成云端训练迁移、安装包打包、GitHub 托管和结项论文/PPT 材料。")

    add_heading(doc, "1.3 技术路线概述", 2)
    add_body(
        doc,
        "项目技术路线可概括为“数据构建—基线训练—分类微调—桌面封装—云端归档”。首先基于 CompCars 数据集筛选汽车外观图像，经过裁剪、清洗和 256×256 标准化处理，形成通用训练数据；随后训练 StyleGAN2-ADA baseline，并根据车型属性构建跑车、轿车、SUV 三类子数据集；在通用 baseline 基础上分别微调三类模型，以稳定满足用户按车型选择生成的功能要求；最后将三类模型集成到 Tkinter 桌面程序中，并通过 PyInstaller 和 Inno Setup 完成 Windows 安装包制作。"
    )

    add_heading(doc, "2 项目需求与总体方案", 1)
    add_heading(doc, "2.1 功能需求分析", 2)
    add_body(
        doc,
        "从结项验收角度看，本项目系统应满足三个层面的需求。第一，模型层面需要具备汽车外观图像生成能力，并能够体现跑车、轿车、SUV 三类结果之间的差异；第二，系统层面需要具有可运行的用户界面，使非模型开发人员也能通过按钮和参数生成图片；第三，交付层面需要形成完整的项目文件、可安装程序和结项说明材料，便于后续继续研究或展示。"
    )
    add_table(
        doc,
        ["需求类别", "具体要求", "实现情况"],
        [
            ["模型能力", "生成汽车外观概念图，支持跑车、轿车、SUV 三类", "通过三类独立微调模型实现"],
            ["交互功能", "车型选择、随机种子、truncation、单张/批量生成", "桌面程序已实现并验证"],
            ["结果保存", "自动保存生成图片并记录参数", "输出至 outputs/app_generated 并生成 metadata.csv"],
            ["部署交付", "形成 Win11 可安装程序", "PyInstaller + Inno Setup 打包完成"],
            ["复现维护", "保留训练脚本、文档、模型和实际数据集", "GitHub 仓库与 Release 资产整理中"],
        ],
    )
    add_caption(doc, "表 1 项目主要功能需求与实现情况")

    add_heading(doc, "2.2 三类模型方案选择", 2)
    add_body(
        doc,
        "在实现“按选择生成三类汽车外观图”这一目标时，项目曾比较两种方案：一种是训练单一条件生成模型，通过类别标签控制输出；另一种是训练三个车型类型的独立微调模型，通过程序加载不同模型实现选择控制。考虑到项目已有通用 StyleGAN2-ADA baseline、训练时间有限、条件模型需要重新构建标签格式并修改训练流程，最终选择“三类数据集 + 三类微调模型”路线。该方案工程复杂度较低，能充分复用已有训练成果，并能较稳定地满足 SRTP 结项演示中的类别选择需求。"
    )

    add_heading(doc, "3 数据集构建与预处理", 1)
    add_heading(doc, "3.1 数据来源与清洗", 2)
    add_body(
        doc,
        "项目使用 CompCars 相关汽车图像数据作为基础数据来源。原始数据包含不同车型、年份、视角和拍摄条件下的汽车图片。为适配汽车外观设计辅助生成场景，项目重点保留外观轮廓较清晰、主体较完整、适合正侧或近似正侧视角学习的图像，并通过清洗脚本进行图像裁剪、方形化和尺寸标准化处理。清洗后得到 49172 张 256×256 汽车外观图像，并进一步生成 StyleGAN2-ADA 所需的 zip 格式训练包。"
    )
    add_body(
        doc,
        "清洗过程的核心目标不是保留所有原始图像，而是尽可能保证训练样本的视觉一致性。对于生成模型而言，输入数据分布越集中，模型越容易学习到稳定的汽车轮廓、车身比例和视角规律。因此项目在早期阶段将精力集中在图像标准化和异常样本剔除上，为后续模型训练提供较干净的数据基础。"
    )

    add_heading(doc, "3.2 三类车型数据集划分", 2)
    add_body(
        doc,
        "为了实现跑车、轿车和 SUV 三类控制生成，项目根据 CompCars 的车型属性信息建立类别映射，将清洗后的图像重新划分为三类训练数据集。映射文件保存为 srtp_type_mapping.json，其中部分原始类型编号被归并到目标类别。例如，原始类型 4 映射为 sedan，类型 2 和 10 映射为 suv，类型 9、11、12 映射为 sports。"
    )
    add_table(
        doc,
        ["目标类别", "英文标识", "样本数量", "用途"],
        [
            ["跑车", "sports", "1519 张", "训练跑车外观生成模型 sports.pkl"],
            ["轿车", "sedan", "5492 张", "训练轿车外观生成模型 sedan.pkl"],
            ["SUV", "suv", "9889 张", "训练 SUV 外观生成模型 suv.pkl"],
        ],
    )
    add_caption(doc, "表 2 三类车型训练数据集规模")
    add_body(
        doc,
        "三类数据集规模存在差异，其中跑车类样本数量明显少于轿车和 SUV。这一差异直接影响训练效果和过拟合风险：跑车模型更容易受到样本数量限制，追加训练时需要更加谨慎；轿车和 SUV 样本数量较多，模型更容易学习到稳定的车身比例与轮廓特征。"
    )

    add_heading(doc, "4 生成模型方法与训练实现", 1)
    add_heading(doc, "4.1 StyleGAN2-ADA 模型", 2)
    add_body(
        doc,
        "StyleGAN2-ADA 是 StyleGAN2 的改进版本，其核心优势在于能够通过自适应数据增强缓解小数据集训练中的判别器过拟合问题。模型通过潜变量 z 生成中间潜空间表示，再经由生成网络逐层合成图像。truncation 参数可控制生成结果的稳定性和多样性：较低 truncation 值通常生成更稳定、更接近训练分布中心的结果；较高 truncation 值会增加变化性，但也可能带来形变或伪影。"
    )
    add_body(
        doc,
        "本项目以 StyleGAN2-ADA PyTorch 版本为基础，并针对新版 PyTorch、CUDA 12.8 和 RTX 5090 环境进行了兼容性处理。桌面程序在加载模型时通过 legacy.load_network_pkl 读取 G_ema 生成器，并根据用户输入的 seed 和 truncation 参数生成图片。"
    )

    add_heading(doc, "4.2 分类型微调策略", 2)
    add_body(
        doc,
        "项目采用分类型微调策略。首先使用清洗后的通用汽车外观数据训练 baseline 模型，使模型学习汽车外观图像的基本分布；随后分别使用跑车、轿车和 SUV 三类数据集从 baseline 权重继续训练，得到三个类型模型。相较于从零开始训练，这种方式能够利用通用模型已经学习到的车辆轮廓、视角和背景规律，减少训练时间并提升小数据集训练的稳定性。"
    )
    add_table(
        doc,
        ["模型文件", "目标车型", "训练数据", "桌面程序加载逻辑"],
        [
            ["sports.pkl", "跑车", "compcars_sports_256_square.zip", "选择“跑车”时加载"],
            ["sedan.pkl", "轿车", "compcars_sedan_256_square.zip", "选择“轿车”时加载"],
            ["suv.pkl", "SUV", "compcars_suv_256_square.zip", "选择“SUV”时加载"],
        ],
    )
    add_caption(doc, "表 3 三类微调模型与程序加载关系")

    add_heading(doc, "5 云端迁移与训练过程", 1)
    add_heading(doc, "5.1 云端环境配置", 2)
    add_body(
        doc,
        "由于本地电脑内存和显存资源有限，且无法长时间保持开机，项目中后期将训练工作整体迁移至云端 GPU 容器。云端实例采用 RTX 5090 32GB 显卡、Ubuntu 22.04、Python 3.12.3、PyTorch 2.7.0+cu128 和 CUDA Runtime 12.8。经验证，PyTorch 能正确识别 RTX 5090，CUDA 可用，矩阵计算测试正常，支持 Blackwell 架构 sm_120。"
    )
    add_table(
        doc,
        ["项目", "配置"],
        [
            ["GPU", "NVIDIA GeForce RTX 5090，32GB 显存"],
            ["CPU", "Intel Xeon Platinum 8470Q，25 核/GPU"],
            ["系统", "Ubuntu 22.04"],
            ["Python", "3.12.3"],
            ["PyTorch", "2.7.0+cu128"],
            ["CUDA Runtime", "12.8"],
            ["项目目录", "/root/autodl-tmp/car_srtp_project"],
        ],
    )
    add_caption(doc, "表 4 云端训练环境配置")

    add_heading(doc, "5.2 训练参数与过程", 2)
    add_body(
        doc,
        "正式训练采用脚本化方式在云端顺序训练三类模型。由于容器环境未安装 tmux，项目使用 nohup 将训练任务放入后台运行，并通过日志文件和 PID 文件监控训练状态。第一轮训练命令为 cloud_train_all_type_models.sh 800 1000 1000 16，表示跑车训练 800 kimg、轿车训练 1000 kimg、SUV 训练 1000 kimg，batch size 为 16。"
    )
    add_body(
        doc,
        "第一轮训练后，三类 fakes 图已经具有基本车辆形状，且三类模型输出之间出现明显差异，但部分图像细节质量和稳定性不足以完全满足结项展示要求。因此项目继续进行追加训练：跑车模型追加 1000 kimg，轿车模型追加 1500 kimg，SUV 模型追加 1500 kimg。追加训练从第一轮最终模型继续，而非回到通用 baseline，既节约训练时间，也保持了类别特征。"
    )
    add_table(
        doc,
        ["阶段", "跑车", "轿车", "SUV", "batch"],
        [
            ["第一轮训练", "800 kimg", "1000 kimg", "1000 kimg", "16"],
            ["追加训练", "+1000 kimg", "+1500 kimg", "+1500 kimg", "16"],
            ["最终样例", "seed 0-47", "seed 0-47", "seed 0-47", "trunc=0.65"],
        ],
    )
    add_caption(doc, "表 5 三类模型训练与追加训练设置")

    add_heading(doc, "6 桌面系统设计与实现", 1)
    add_heading(doc, "6.1 系统架构", 2)
    add_body(
        doc,
        "桌面程序采用 Python Tkinter/ttk 构建图形界面，后端使用 PyTorch 加载 StyleGAN2-ADA 生成器模型，Pillow 负责图像显示和保存。系统启动时检查依赖环境并初始化生成后端；用户选择车型后，程序查找对应的 pkl 模型文件，并根据 seed 与 truncation 参数执行图像生成。生成结果会显示在界面预览区，并自动保存到 outputs/app_generated 目录，同时将生成时间、车型、seed、truncation 和图片路径记录到 metadata.csv。"
    )
    add_table(
        doc,
        ["模块", "技术/文件", "功能"],
        [
            ["界面层", "Tkinter / ttk", "车型选择、参数输入、按钮交互、图片预览"],
            ["模型层", "StyleGAN2-ADA + PyTorch", "加载 G_ema 生成器并执行推理"],
            ["图像层", "Pillow / NumPy", "图像数组转换、显示与保存"],
            ["记录层", "metadata.csv", "记录生成参数和输出路径"],
            ["部署层", "PyInstaller + Inno Setup", "生成 Win11 可运行安装包"],
        ],
    )
    add_caption(doc, "表 6 桌面系统模块划分")

    add_heading(doc, "6.2 交互界面与功能", 2)
    add_body(
        doc,
        "桌面程序界面由顶部标题区、左侧控制面板和右侧图片预览区组成。左侧控制面板提供跑车、轿车、SUV 三类单选按钮，并显示当前类别的简要说明；用户可输入或随机生成 seed，通过滑块调节 truncation 稳定度，并设置批量生成数量。程序提供“生成单张”“批量生成”“打开输出目录”“另存当前图片”等按钮，能够满足结项演示中快速切换车型并展示生成结果的需求。"
    )
    add_body(
        doc,
        "在工程实现中，程序对源码运行和 PyInstaller 打包运行分别处理路径问题。源码运行时优先从 app/models 或 outputs/final_models 读取模型；打包运行时还兼容 _internal/app/models 目录。该路径兼容处理解决了打包后找不到 legacy、click 和三类 pkl 模型的问题，使安装版程序能够在其他 Win11 电脑上正常运行并生成图片。"
    )

    add_heading(doc, "6.3 打包与部署", 2)
    add_body(
        doc,
        "项目先使用 PyInstaller 将 Python 桌面程序、PyTorch 依赖、StyleGAN2-ADA 运行代码和模型文件打包为 onedir 程序目录，再使用 Inno Setup 制作 Windows 安装包。打包过程中解决了多个工程问题：包括 PyInstaller 未自动收集 click、legacy 等依赖，StyleGAN2-ADA 代码路径在 _internal 目录下无法被发现，三类模型文件未复制到程序查找路径，以及 Inno Setup 中文语言包缺失等。最终安装包在本机安装后能够正常启动并生成三类汽车图像。"
    )
    add_body(
        doc,
        "由于最终安装包体积超过 GitHub Release 单文件上传的稳定范围，项目进一步将 Inno Setup 配置为分卷输出，生成一个启动 exe 和多个 bin 数据分卷。该方式便于在 GitHub Release 中保存大型安装程序，也便于后续在云端重新下载和复现系统。"
    )

    add_heading(doc, "7 实验结果与结项成果", 1)
    add_heading(doc, "7.1 生成效果评价", 2)
    add_body(
        doc,
        "追加训练完成后，三类模型均能够生成具有车辆基本轮廓的汽车外观图像。跑车模型生成结果倾向于较低车身和运动姿态，轿车模型生成结果更接近日常乘用车比例，SUV 模型则更容易呈现较高车身和更厚重的体量感。虽然生成图像在细节真实性、车灯结构、轮胎边缘和背景稳定性方面仍存在不足，但从 SRTP 结项演示角度看，系统已经能够体现三类车型之间的可辨别差异，并能通过随机 seed 生成多张候选概念图。"
    )
    add_body(
        doc,
        "项目采用人工筛选方式对最终生成图进行展示性评估。当前每类均能够筛选出 10 张以上可展示图片，满足结项展示和 PPT 汇报需要。由于本项目主要目标是开发可运行的辅助系统，而非达到商业级汽车渲染质量，因此评价重点放在系统完整性、类别可控性和生成结果可展示性上。"
    )

    add_heading(doc, "7.2 最终交付成果", 2)
    add_table(
        doc,
        ["成果类别", "具体成果", "说明"],
        [
            ["模型成果", "sports.pkl、sedan.pkl、suv.pkl", "三类 StyleGAN2-ADA 微调模型"],
            ["数据成果", "三类 type_datasets zip", "项目实际用于三类训练的数据集"],
            ["程序成果", "car_design_assistant.py", "桌面程序源码"],
            ["安装成果", "Inno Setup 分卷安装包", "安装后应用可正常生图"],
            ["文档成果", "README、cloud_env_config、结项报告", "记录项目过程和复现方法"],
            ["云端成果", "GitHub 仓库与 Release 资产", "便于后续继续开发和云端保存"],
        ],
    )
    add_caption(doc, "表 7 项目结项交付成果")

    add_heading(doc, "7.3 GitHub 托管与项目归档", 2)
    add_body(
        doc,
        "为减少本地电脑存储压力并便于后续继续推进，项目已上传至 GitHub 仓库 jinglangqiu-hub/car-design-assistant-srtp。仓库中保存源码、训练脚本、打包配置和文档；模型、安装包分卷、三类训练数据集和样例图作为 Release 资产上传。该方式避免将大文件直接写入 Git 历史，同时保留项目继续训练、重新打包和展示部署所需的核心资源。"
    )

    add_heading(doc, "8 问题分析、改进方向与总结", 1)
    add_heading(doc, "8.1 存在问题", 2)
    add_body(
        doc,
        "尽管项目完成了从数据处理、云端训练到桌面系统部署的完整流程，但仍存在若干不足。首先，当前模型分辨率为 256×256，适合概念图展示，但无法提供高分辨率细节；其次，三类模型采用独立微调方式实现类别控制，工程上稳定但模型数量较多，后续若类别扩展会增加维护成本；再次，跑车类样本数量较少，模型在部分随机种子下仍可能出现车身形变或细节模糊；最后，当前系统只支持车型类别、seed 和 truncation 控制，还不能通过文字描述、草图或局部编辑进行更精细的设计控制。"
    )

    add_heading(doc, "8.2 后续改进方向", 2)
    add_bullet(doc, "构建 512×512 或更高分辨率数据集，提升生成图像清晰度和细节表现。")
    add_bullet(doc, "尝试单一条件生成模型，将车型标签作为条件输入，减少模型数量并提升扩展性。")
    add_bullet(doc, "引入 CLIP、文本提示或草图控制，实现更符合设计需求的可控生成。")
    add_bullet(doc, "增加生成结果图库管理、收藏、评分和批量导出功能。")
    add_bullet(doc, "扩大跑车类样本规模，改善少样本类别生成稳定性。")

    add_heading(doc, "8.3 总结", 2)
    add_body(
        doc,
        "本项目围绕“汽车外观设计辅助系统”这一 SRTP 主题，完成了从数据集构建、生成模型训练、云端环境迁移、桌面程序开发、安装包制作到 GitHub 归档的完整工程闭环。项目最终实现了按跑车、轿车、SUV 三类选择生成汽车外观概念图的功能，并形成可运行的 Windows 桌面程序和可安装交付包。虽然生成图像质量距离真实汽车设计生产流程仍有差距，但项目验证了生成式模型在汽车外观早期概念辅助中的可行性，也为后续引入更高分辨率训练、条件生成和多模态控制奠定了基础。"
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Goodfellow I., Pouget-Abadie J., Mirza M., et al. Generative Adversarial Nets. Advances in Neural Information Processing Systems, 2014.",
        "[2] Karras T., Laine S., Aila T. A Style-Based Generator Architecture for Generative Adversarial Networks. CVPR, 2019.",
        "[3] Karras T., Laine S., Aittala M., et al. Analyzing and Improving the Image Quality of StyleGAN. CVPR, 2020.",
        "[4] Karras T., Aittala M., Hellsten J., et al. Training Generative Adversarial Networks with Limited Data. NeurIPS, 2020.",
        "[5] Yang L., Luo P., Change Loy C., Tang X. A Large-Scale Car Dataset for Fine-Grained Categorization and Verification. CVPR, 2015.",
        "[6] NVIDIA. StyleGAN2-ADA PyTorch official repository and documentation.",
        "[7] GitHub Docs. About releases and managing large files.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        set_run_font(r, size=10.5)

    add_heading(doc, "附录 A 项目关键路径", 1)
    add_table(
        doc,
        ["路径", "说明"],
        [
            ["app/car_design_assistant.py", "桌面程序主入口"],
            ["app/models/", "桌面程序运行使用的三类模型"],
            ["data/type_datasets/", "三类训练数据集"],
            ["scripts/cloud_train_all_type_models.sh", "第一轮云端正式训练脚本"],
            ["scripts/cloud_continue_all_type_models.sh", "追加训练脚本"],
            ["scripts/build_pyinstaller_app.ps1", "PyInstaller 打包脚本"],
            ["scripts/build_inno_installer.ps1", "Inno Setup 安装包脚本"],
            ["packaging/CarDesignAssistant.iss", "安装包配置文件"],
        ],
    )

    add_heading(doc, "附录 B 项目开发历史摘要", 1)
    for item in [
        "完成本地数据清洗与通用 StyleGAN2-ADA baseline 训练。",
        "迁移至 RTX 5090 云端实例，验证 PyTorch 2.7.0+cu128 和 CUDA 12.8 环境。",
        "构建跑车、轿车、SUV 三类数据集，并完成三类冒烟训练。",
        "执行第一轮正式训练，获得三类初步可区分模型。",
        "根据生成质量不足的问题追加训练，提升可展示样例数量。",
        "开发 Tkinter 桌面程序，实现三类选择生成。",
        "解决 PyInstaller 依赖、模型路径和 StyleGAN2-ADA 打包兼容问题。",
        "使用 Inno Setup 生成安装包，并验证安装后可正常生图。",
        "整理项目目录，删除明确多余的构建缓存和训练中间产物。",
        "上传源码至 GitHub，并整理 Release 资产用于云端归档。",
    ]:
        add_numbered(doc, item)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_report()
